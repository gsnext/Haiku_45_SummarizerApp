"""Text extraction utilities for multiple document formats."""
import io
from typing import Literal
from pathlib import Path
import requests
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
from ..errors import FileFormatError, ExtractionError, FileSizeError, URLFetchError
from ..logger import logger
from ..config import settings


async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file content."""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        logger.info("Successfully extracted text from PDF")
        return text
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise ExtractionError(f"Failed to extract text from PDF: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file content."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        logger.info("Successfully extracted text from DOCX")
        return text
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise ExtractionError(f"Failed to extract text from DOCX: {str(e)}")


async def extract_text_from_url(url: str) -> str:
    """Fetch and extract text from a URL."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = " ".join(chunk for chunk in chunks if chunk)

        logger.info(f"Successfully extracted text from URL: {url}")
        return text
    except requests.RequestException as e:
        logger.error(f"URL fetch failed: {str(e)}")
        raise URLFetchError(f"Failed to fetch content from URL: {str(e)}")
    except Exception as e:
        logger.error(f"URL text extraction failed: {str(e)}")
        raise ExtractionError(f"Failed to extract text from URL: {str(e)}")


async def extract_text(content: str | bytes, format_type: Literal["txt", "pdf", "docx", "url"]) -> str:
    """
    Extract text from various document formats.

    Args:
        content: File content or URL string
        format_type: Format of the content (txt, pdf, docx, url)

    Returns:
        Extracted text content
    """
    try:
        if format_type == "txt":
            if isinstance(content, bytes):
                return content.decode("utf-8", errors="ignore")
            return content

        elif format_type == "pdf":
            if isinstance(content, str):
                content = content.encode()
            return await extract_text_from_pdf(content)

        elif format_type == "docx":
            if isinstance(content, str):
                content = content.encode()
            return await extract_text_from_docx(content)

        elif format_type == "url":
            if isinstance(content, bytes):
                content = content.decode("utf-8")
            return await extract_text_from_url(content)

        else:
            raise FileFormatError(f"Unsupported format: {format_type}")

    except Exception as e:
        logger.error(f"Text extraction error: {str(e)}")
        raise


def validate_file_size(file_size: int) -> None:
    """Validate that file size does not exceed maximum allowed."""
    if file_size > settings.MAX_FILE_SIZE:
        raise FileSizeError(
            f"File size exceeds maximum allowed size of {settings.MAX_FILE_SIZE / (1024*1024):.1f}MB"
        )


def validate_format(format_type: str) -> None:
    """Validate that format is supported."""
    if format_type.lower() not in settings.ALLOWED_FORMATS:
        raise FileFormatError(
            f"Unsupported format: {format_type}. Allowed formats: {', '.join(settings.ALLOWED_FORMATS)}"
        )
